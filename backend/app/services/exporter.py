"""Export service: XLS for Yandex Direct Commander and Markdown strategy."""
from __future__ import annotations

import io
import zipfile
from typing import TYPE_CHECKING

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

if TYPE_CHECKING:
    from sqlalchemy.orm import Session


# ── Column layout (66 cols) matching the official Direct Commander template ──
# Row 10 — main headers (only named cols, others blank)
_H1: dict[int, str] = {
    1: "Доп. объявление группы", 2: "Тип объявления", 3: "ID группы",
    4: "Название группы", 5: "Номер группы", 6: "ID фразы",
    7: "Фраза (с минус-словами)", 8: "ID объявления",
    9: "Заголовок 1", 10: "Заголовок 2", 11: "Текст",
    12: "Длина",         # spans 12-14 (visual merge only)
    15: "Комбинаторика", # spans 15-46
    47: "Ссылка", 48: "Отображаемая ссылка", 49: "Регион",
    50: "Организация Яндекс Бизнеса", 51: "Ставка", 52: "Ставка в сетях",
    53: "Статус объявления ", 54: "Статус фразы",
    55: "Заголовки быстрых ссылок", 56: "Описания быстрых ссылок",
    57: "Адреса быстрых ссылок", 58: "Параметр 1", 59: "Параметр 2",
    60: "Метки", 61: "Изображение", 62: "Креатив",
    63: "Статус модерации креатива", 64: "Уточнения",
    65: "Минус-фразы на группу", 66: "Возрастные ограничения",
}

# Row 11 — sub-headers under «Длина» (12-14) and «Комбинаторика» (15-46)
_H2: dict[int, str] = {
    12: "заголовок 1", 13: "заголовок 2", 14: "текст",
    15: "Заголовок 1", 16: "Заголовок 2", 17: "Заголовок 3",
    18: "Заголовок 4", 19: "Заголовок 5", 20: "Заголовок 6",
    21: "Заголовок 7", 22: "Текст 1", 23: "Текст 2", 24: "Текст 3",
    25: "Длина заголовка 1", 26: "Длина заголовка 2", 27: "Длина заголовка 3",
    28: "Длина заголовка 4", 29: "Длина заголовка 5", 30: "Длина заголовка 6",
    31: "Длина заголовка 7", 32: "Длина текста 1", 33: "Длина текста 2",
    34: "Длина текста 3", 35: "Изображение 1", 36: "Изображение 2",
    37: "Изображение 3", 38: "Изображение 4", 39: "Изображение 5",
    40: "Вертикальное видео 1", 41: "Вертикальное видео 2",
    42: "Квадратное видео 1", 43: "Квадратное видео 2",
    44: "Горизонтальное видео 1", 45: "Горизонтальное видео 2",
    46: "Причины отклонения",
}


def _format_geo(geo) -> str:
    if not geo:
        return ""
    if isinstance(geo, str):
        return geo
    if isinstance(geo, list):
        parts = [g.get("name", str(g)) if isinstance(g, dict) else str(g) for g in geo]
        return ", ".join(parts)
    if isinstance(geo, dict):
        return geo.get("name", str(geo))
    return str(geo)


def _encode_keyword(phrase: str, match_type: str) -> str:
    if match_type == "exact":
        return f'"{phrase}"'
    return phrase


def _build_campaign_xls(campaign, groups, ads_by_group, keywords_by_group,
                         neg_kws, project_url: str) -> bytes:
    """Build one .xlsx file matching the official Yandex Direct Commander EPC template."""
    wb = Workbook()
    wb.remove(wb.active)

    ws = wb.create_sheet("Тексты")

    # ── Campaign header block (rows 6-9) ────────────────────────────────────
    ws.cell(6, 1, "Предложение текстовых блоков для кампании")
    ws.cell(7, 4, "Тип кампании:")
    ws.cell(7, 5, "Единая перфоманс-кампания")
    ws.cell(8, 4, "№ заказа:")
    ws.cell(8, 5, "")
    ws.cell(8, 7, "Валюта:")
    ws.cell(8, 8, "RUB")
    ws.cell(9, 4, "Минус-фразы на кампанию:")
    minus_parts = []
    for nk in neg_kws:
        p = nk.phrase.strip()
        if p:
            minus_parts.append(f"-{p}" if not p.startswith("-") else p)
    ws.cell(9, 5, " ".join(minus_parts))

    # ── Column headers (rows 10-11) ─────────────────────────────────────────
    fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
    bold = Font(bold=True)
    for col, val in _H1.items():
        cell = ws.cell(10, col, val)
        cell.font = bold
        cell.fill = fill
    for col, val in _H2.items():
        cell = ws.cell(11, col, val)
        cell.font = bold
        cell.fill = fill

    ws.freeze_panes = "A12"

    geo_str = _format_geo(campaign.geo)
    current_row = 12

    # ── Data rows ────────────────────────────────────────────────────────────
    for group_idx, group in enumerate(groups, 1):
        gid = str(group.id)
        keywords = keywords_by_group.get(gid, [])
        ads = ads_by_group.get(gid, [])

        if not ads and not keywords:
            continue

        # Prefer "ready" ads; fall back to any ads
        ready = [a for a in ads if a.status.value == "ready"] or ads
        first_ad = ready[0] if ready else None
        extra_ads = ready[1:] if len(ready) > 1 else []

        h1  = (first_ad.headline1 or "") if first_ad else ""
        h2  = (first_ad.headline2 or "") if first_ad else ""
        txt = (first_ad.text or "")       if first_ad else ""
        url = (first_ad.display_url or project_url) if first_ad else project_url

        def _row(r, marker, phrase, rh1, rh2, rtxt, rurl):
            ws.cell(r, 1, marker)
            ws.cell(r, 2, "Текстово-графическое")
            ws.cell(r, 4, group.name)
            ws.cell(r, 5, group_idx)
            ws.cell(r, 7, phrase)
            ws.cell(r, 9,  rh1)
            ws.cell(r, 10, rh2)
            ws.cell(r, 11, rtxt)
            ws.cell(r, 12, len(rh1))
            ws.cell(r, 13, len(rh2))
            ws.cell(r, 14, len(rtxt))
            for c in range(15, 47):   # combinatorial slots — empty for text ads
                ws.cell(r, c, "")
            ws.cell(r, 47, rurl)
            ws.cell(r, 49, geo_str)

        # One "-" row per keyword (all share the primary ad creative)
        phrases = [_encode_keyword(kw.phrase, kw.match_type) for kw in keywords]
        if not phrases:
            phrases = ["---autotargeting"]

        for phrase in phrases:
            _row(current_row, "-", phrase, h1, h2, txt, url)
            current_row += 1

        # Additional ad variants → "+" rows (no keyword)
        for extra in extra_ads:
            eh1  = extra.headline1 or ""
            eh2  = extra.headline2 or ""
            etxt = extra.text or ""
            eurl = extra.display_url or project_url
            _row(current_row, "+", "", eh1, eh2, etxt, eurl)
            current_row += 1

    # ── Reference sheets (minimal stubs) ────────────────────────────────────
    ws_r = wb.create_sheet("Регионы")
    ws_r.cell(3, 2, "Регионы")

    ws_d = wb.create_sheet("Словарь значений полей")
    ws_d.cell(2, 2, "Тип кампании")
    ws_d.cell(3, 2, "Единая перфоманс-кампания")
    ws_d.cell(5, 2, "Тип объявления")
    ws_d.cell(6, 2, "Текстово-графическое")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def export_direct_xls(project_id, db: "Session") -> tuple[bytes, str]:
    """Export Direct campaigns as Commander XLS files.

    Returns ``(data, ext)`` where *ext* is ``"xlsx"`` for a single campaign
    or ``"zip"`` when multiple campaigns are packed into an archive.
    """
    import uuid

    from sqlalchemy import select

    from app.models.direct import Ad, AdGroup, Campaign, Keyword, NegativeKeyword
    from app.models.project import Project

    if not isinstance(project_id, uuid.UUID):
        project_id = uuid.UUID(str(project_id))

    project = db.get(Project, project_id)
    if not project:
        raise ValueError("Project not found")

    campaigns = db.scalars(
        select(Campaign).where(Campaign.project_id == project_id).order_by(Campaign.priority)
    ).all()
    if not campaigns:
        raise ValueError("No campaigns to export")

    campaign_ids = [c.id for c in campaigns]

    groups_all = db.scalars(
        select(AdGroup).where(AdGroup.campaign_id.in_(campaign_ids))
    ).all()
    group_ids = [g.id for g in groups_all]

    keywords_all = db.scalars(
        select(Keyword).where(Keyword.ad_group_id.in_(group_ids))
    ).all() if group_ids else []

    ads_all = db.scalars(
        select(Ad).where(Ad.ad_group_id.in_(group_ids))
    ).all() if group_ids else []

    neg_kws_all = db.scalars(
        select(NegativeKeyword).where(NegativeKeyword.project_id == project_id)
    ).all()

    # Build lookup maps
    groups_by_campaign: dict[str, list] = {}
    for g in groups_all:
        groups_by_campaign.setdefault(str(g.campaign_id), []).append(g)

    ads_by_group: dict[str, list] = {}
    for a in ads_all:
        ads_by_group.setdefault(str(a.ad_group_id), []).append(a)

    keywords_by_group: dict[str, list] = {}
    for kw in keywords_all:
        keywords_by_group.setdefault(str(kw.ad_group_id), []).append(kw)

    # Project-level negatives go into every campaign
    project_neg_kws = [nk for nk in neg_kws_all if not nk.campaign_id]

    project_url = project.url or ""

    xls_files: list[tuple[str, bytes]] = []
    for campaign in campaigns:
        cid = str(campaign.id)
        groups = groups_by_campaign.get(cid, [])
        camp_neg_kws = [nk for nk in neg_kws_all if nk.campaign_id == campaign.id] + project_neg_kws

        data = _build_campaign_xls(
            campaign, groups, ads_by_group, keywords_by_group, camp_neg_kws, project_url
        )
        safe_name = campaign.name.replace("/", "-").replace("\\", "-")[:80]
        xls_files.append((f"{safe_name}.xlsx", data))

    if len(xls_files) == 1:
        return xls_files[0][1], "xlsx"

    # Multiple campaigns → ZIP archive
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in xls_files:
            zf.writestr(name, data)
    zip_buf.seek(0)
    return zip_buf.read(), "zip"


def export_strategy_md(project_id, db: "Session") -> str:
    """Export Direct strategy as Markdown."""
    import uuid

    from sqlalchemy import select

    from app.models.brief import Brief
    from app.models.direct import AdGroup, Campaign, Keyword
    from app.models.project import Project

    if not isinstance(project_id, uuid.UUID):
        project_id = uuid.UUID(str(project_id))

    project = db.get(Project, project_id)
    if not project:
        raise ValueError("Project not found")

    brief = db.scalar(select(Brief).where(Brief.project_id == project_id))
    campaigns = db.scalars(
        select(Campaign).where(Campaign.project_id == project_id).order_by(Campaign.priority)
    ).all()

    lines = []
    lines.append(f"# Стратегия Яндекс Директ — {project.name}")
    lines.append(f"**Клиент:** {project.client_name}  ")
    lines.append(f"**Сайт:** {project.url}  ")
    lines.append("")

    if brief:
        lines.append("## Бриф")
        if brief.niche:
            lines.append(f"- **Ниша:** {brief.niche}")
        if brief.products:
            lines.append(f"- **Продукты/услуги:** {brief.products}")
        if brief.usp:
            lines.append(f"- **УТП:** {brief.usp}")
        if brief.campaign_goal:
            lines.append(f"- **Цель кампании:** {brief.campaign_goal}")
        if brief.ad_geo:
            lines.append(f"- **Гео:** {', '.join(brief.ad_geo)}")
        if brief.monthly_budget:
            lines.append(f"- **Бюджет:** {brief.monthly_budget} ₽/мес")
        lines.append("")

    lines.append("## Кампании")
    for c in campaigns:
        lines.append(f"### {c.name}")
        if c.type:
            lines.append(f"**Тип:** {c.type}")
        if c.budget_monthly:
            lines.append(f"**Бюджет:** {c.budget_monthly} ₽/мес")
        if c.strategy_text:
            lines.append("")
            lines.append(c.strategy_text)
        lines.append("")

        # Groups + keywords
        groups = db.scalars(select(AdGroup).where(AdGroup.campaign_id == c.id)).all()
        for g in groups:
            lines.append(f"#### Группа: {g.name}")
            keywords = db.scalars(
                select(Keyword).where(Keyword.ad_group_id == g.id).order_by(Keyword.frequency.desc().nullslast())
            ).all()
            if keywords:
                lines.append("")
                lines.append("| Фраза | Частота | Температура |")
                lines.append("|-------|---------|-------------|")
                for kw in keywords[:20]:
                    freq = str(kw.frequency) if kw.frequency else "—"
                    temp = kw.temperature.value if kw.temperature else "—"
                    lines.append(f"| {kw.phrase} | {freq} | {temp} |")
            lines.append("")

    return "\n".join(lines)


def validate_export(project_id, db: "Session") -> dict:
    """Pre-export validation summary."""
    import uuid

    from sqlalchemy import func, select

    from app.models.direct import Ad, AdGroup, Campaign, Keyword, NegativeKeyword

    if not isinstance(project_id, uuid.UUID):
        project_id = uuid.UUID(str(project_id))

    campaigns = db.scalars(select(Campaign).where(Campaign.project_id == project_id)).all()
    campaign_ids = [c.id for c in campaigns]
    groups = db.scalars(select(AdGroup).where(AdGroup.campaign_id.in_(campaign_ids))).all() if campaign_ids else []
    group_ids = [g.id for g in groups]
    keywords = db.scalars(select(Keyword).where(Keyword.ad_group_id.in_(group_ids))).all() if group_ids else []
    ads = db.scalars(select(Ad).where(Ad.ad_group_id.in_(group_ids))).all() if group_ids else []
    neg_kws = db.scalars(select(NegativeKeyword).where(NegativeKeyword.project_id == project_id)).all()

    warnings = []
    invalid_ads = []
    for ad in ads:
        issues = []
        if len(ad.headline1 or "") > 56:
            issues.append("headline1 > 56 chars")
        if len(ad.headline2 or "") > 30:
            issues.append("headline2 > 30 chars")
        if len(ad.headline3 or "") > 30:
            issues.append("headline3 > 30 chars")
        if len(ad.text or "") > 81:
            issues.append("text > 81 chars")
        if issues:
            invalid_ads.append({"ad_id": str(ad.id), "issues": issues})

    if invalid_ads:
        warnings.append(f"{len(invalid_ads)} объявлений превышают лимиты символов")

    if not keywords:
        warnings.append("Семантическое ядро пустое")

    if not ads:
        warnings.append("Объявления не созданы")

    return {
        "campaigns_count": len(campaigns),
        "groups_count": len(groups),
        "ads_count": len(ads),
        "keywords_count": len(keywords),
        "negative_keywords_count": len(neg_kws),
        "warnings": warnings,
        "invalid_ads": invalid_ads,
        "ready": len(warnings) == 0,
    }


# ─────────────────────────────────────────────────────────────────────────────
# DOCX: ТЗ для копирайтера
# ─────────────────────────────────────────────────────────────────────────────

def export_copywriter_docx(project_id, db: "Session") -> bytes:
    """Generate a DOCX copywriting brief for the project."""
    import uuid

    from docx import Document
    from docx.shared import Pt, RGBColor
    from sqlalchemy import select

    from app.models.brief import Brief
    from app.models.crawl import CrawlSession, CrawlStatus, Page
    from app.models.direct import AdGroup, Campaign, Keyword, NegativeKeyword
    from app.models.project import Project
    from app.models.seo import SeoPageMeta

    if not isinstance(project_id, uuid.UUID):
        project_id = uuid.UUID(str(project_id))

    project = db.get(Project, project_id)
    if not project:
        raise ValueError("Project not found")

    brief = db.scalar(select(Brief).where(Brief.project_id == project_id))

    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"ТЗ копирайтеру — {project.name}", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # primary blue

    doc.add_paragraph(f"Клиент: {project.client_name}")
    doc.add_paragraph(f"Сайт: {project.url}")
    doc.add_paragraph("")

    # ── Brief ──────────────────────────────────────────────────────────────────
    if brief:
        doc.add_heading("О бизнесе", level=1)
        if brief.niche:
            p = doc.add_paragraph()
            p.add_run("Ниша: ").bold = True
            p.add_run(brief.niche)
        if brief.products:
            p = doc.add_paragraph()
            p.add_run("Продукты/услуги: ").bold = True
            p.add_run(brief.products)
        if brief.usp:
            p = doc.add_paragraph()
            p.add_run("УТП: ").bold = True
            p.add_run(brief.usp)
        if brief.target_audience:
            p = doc.add_paragraph()
            p.add_run("ЦА: ").bold = True
            p.add_run(brief.target_audience)
        if brief.pains:
            p = doc.add_paragraph()
            p.add_run("Боли клиентов: ").bold = True
            p.add_run(brief.pains)
        if brief.restrictions:
            p = doc.add_paragraph()
            p.add_run("Ограничения: ").bold = True
            p.add_run(brief.restrictions)
        doc.add_paragraph("")

    # ── Pages with SEO recommendations ────────────────────────────────────────
    crawl = db.scalar(
        select(CrawlSession)
        .where(CrawlSession.project_id == project_id, CrawlSession.status == CrawlStatus.DONE)
        .order_by(CrawlSession.finished_at.desc())
    )
    if crawl:
        seo_pages = db.scalars(
            select(SeoPageMeta).where(SeoPageMeta.crawl_session_id == crawl.id)
        ).all()
        if seo_pages:
            doc.add_heading("Рекомендации по страницам", level=1)
            doc.add_paragraph(
                "Для каждой страницы указаны текущие мета-теги и рекомендуемые. "
                "Напишите тексты с учётом УТП и ключевых запросов."
            )
            doc.add_paragraph("")

            for page in seo_pages[:50]:  # limit to 50 pages
                doc.add_heading(page.page_url, level=2)
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Table Grid"
                hdr = tbl.rows[0].cells
                hdr[0].text = "Поле"
                hdr[1].text = "Текущее"
                hdr[2].text = "Рекомендуемое"
                for field, current, rec in [
                    ("Title", page.current_title or "—", page.rec_title or ""),
                    ("Description", page.current_description or "—", page.rec_description or ""),
                    ("OG Title", page.og_title or "—", page.rec_og_title or ""),
                    ("OG Description", page.og_description or "—", page.rec_og_description or ""),
                ]:
                    row = tbl.add_row().cells
                    row[0].text = field
                    row[1].text = current
                    row[2].text = rec
                doc.add_paragraph("")

    # ── Keywords per group ─────────────────────────────────────────────────────
    campaigns = db.scalars(
        select(Campaign).where(Campaign.project_id == project_id).order_by(Campaign.priority)
    ).all()
    if campaigns:
        doc.add_heading("Семантика по группам", level=1)
        doc.add_paragraph(
            "Ключевые запросы, под которые пишутся тексты объявлений и страниц."
        )
        for campaign in campaigns:
            doc.add_heading(campaign.name, level=2)
            groups = db.scalars(
                select(AdGroup).where(AdGroup.campaign_id == campaign.id)
            ).all()
            for group in groups:
                doc.add_heading(f"Группа: {group.name}", level=3)
                keywords = db.scalars(
                    select(Keyword).where(Keyword.ad_group_id == group.id)
                    .order_by(Keyword.frequency.desc().nullslast())
                ).all()
                if keywords:
                    for kw in keywords:
                        freq = f" ({kw.frequency:,})" if kw.frequency else ""
                        doc.add_paragraph(f"• {kw.phrase}{freq}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# HTML: Стратегия для печати / PDF
# ─────────────────────────────────────────────────────────────────────────────

def _get_print_css(primary_color: str = "#1e40af") -> str:
    return f"""
body {{ font-family: 'Helvetica Neue', Arial, sans-serif; max-width: 860px;
       margin: 0 auto; padding: 40px 20px; color: #111; line-height: 1.6; }}
h1 {{ color: {primary_color}; border-bottom: 2px solid {primary_color}; padding-bottom: 8px; }}
h2 {{ color: {primary_color}; margin-top: 32px; }}
h3 {{ color: #374151; }}
h4 {{ color: #6b7280; }}
table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
th, td {{ border: 1px solid #d1d5db; padding: 6px 10px; text-align: left; font-size: 13px; }}
th {{ background: #eff6ff; font-weight: 600; }}
code {{ background: #f3f4f6; padding: 1px 4px; border-radius: 3px; font-size: 12px; }}
blockquote {{ border-left: 4px solid #93c5fd; padding-left: 12px; color: #374151; margin: 12px 0; }}
.cover {{ text-align: center; padding: 60px 20px; border-bottom: 2px solid {primary_color}; margin-bottom: 40px; }}
.cover h1 {{ border: none; font-size: 2em; }}
.cover .subtitle {{ color: #6b7280; margin-top: 8px; }}
.agency-logo {{ max-height: 60px; margin-bottom: 20px; }}
@media print {{
  @page {{ margin: 20mm; }}
  h1, h2 {{ page-break-after: avoid; }}
  table {{ page-break-inside: avoid; }}
  .cover {{ page-break-after: always; }}
}}
"""


def export_strategy_html(project_id, db: "Session") -> str:
    """Convert the Markdown strategy to a print-ready HTML page with white label."""
    import markdown as md_lib

    from app.services.settings_service import get_setting

    agency_name = get_setting("white_label_agency_name", db) or "SEODirect Tool"
    logo_url = get_setting("white_label_logo_url", db) or ""
    primary_color = get_setting("white_label_primary_color", db) or "#1e40af"

    import uuid as _uuid

    from sqlalchemy import select

    from app.models.project import Project
    if not isinstance(project_id, _uuid.UUID):
        project_id = _uuid.UUID(str(project_id))
    project = db.get(Project, project_id)
    project_name = project.name if project else ""
    client_name = project.client_name if project else ""

    md_text = export_strategy_md(project_id, db)
    body = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])

    logo_html = f'<img src="{logo_url}" alt="{agency_name}" class="agency-logo">' if logo_url else ""

    from datetime import date
    today = date.today().strftime("%d.%m.%Y")

    cover = f"""<div class="cover">
  {logo_html}
  <h1>Стратегия Яндекс Директ</h1>
  <p class="subtitle"><strong>{project_name}</strong> · {client_name}</p>
  <p class="subtitle">Подготовлено: {agency_name} · {today}</p>
  <p style="margin-top:16px;font-size:12px;color:#9ca3af">Откройте Ctrl+P → Сохранить как PDF</p>
</div>"""

    css = _get_print_css(primary_color)
    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Стратегия — {project_name}</title>
  <style>{css}</style>
</head>
<body>
{cover}
{body}
<hr style="margin-top:40px;border-color:#e5e7eb">
<p style="font-size:11px;color:#9ca3af;text-align:right">
  Подготовлено {agency_name} · {today}
</p>
</body>
</html>"""


# ─────────────────────────────────────────────────────────────────────────────
# XLSX: Медиаплан
# ─────────────────────────────────────────────────────────────────────────────

def export_mediaplan_xlsx(project_id, db: "Session") -> bytes:
    """Export mediaplan rows as XLSX."""
    import uuid

    from sqlalchemy import select

    from app.models.mediaplan import MediaPlan
    from app.models.project import Project

    if not isinstance(project_id, uuid.UUID):
        project_id = uuid.UUID(str(project_id))

    project = db.get(Project, project_id)
    if not project:
        raise ValueError("Project not found")

    plan = db.scalar(select(MediaPlan).where(MediaPlan.project_id == project_id))
    rows: list[dict] = (plan.rows or []) if plan else []

    wb = Workbook()
    ws = wb.active
    ws.title = "Медиаплан"

    headers = ["Месяц", "% бюджета", "Бюджет (₽)", "Прогноз кликов", "Прогноз заявок", "CPC (₽)", "CPA (₽)"]
    bold = Font(bold=True)
    fill = PatternFill(fill_type="solid", fgColor="1E40AF")
    white_bold = Font(bold=True, color="FFFFFF")

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = white_bold
        cell.fill = fill

    month_names = {
        1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель",
        5: "Май", 6: "Июнь", 7: "Июль", 8: "Август",
        9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь",
    }

    total_budget = sum(r.get("budget") or 0 for r in rows)
    total_clicks = sum(r.get("forecast_clicks") or 0 for r in rows)
    total_leads = sum(r.get("forecast_leads") or 0 for r in rows)

    for row_idx, r in enumerate(rows, 2):
        budget = r.get("budget") or 0
        clicks = r.get("forecast_clicks") or None
        leads = r.get("forecast_leads") or None
        pct = r.get("pct") or (round(budget / total_budget * 100, 1) if total_budget else 0)
        cpc = round(budget / clicks) if budget and clicks else ""
        cpa = round(budget / leads) if budget and leads else ""
        month_num = r.get("month", row_idx - 1)
        month_name = r.get("month_name") or month_names.get(month_num, str(month_num))
        ws.cell(row=row_idx, column=1, value=month_name)
        ws.cell(row=row_idx, column=2, value=f"{pct}%")
        ws.cell(row=row_idx, column=3, value=budget)
        ws.cell(row=row_idx, column=4, value=clicks or "")
        ws.cell(row=row_idx, column=5, value=leads or "")
        ws.cell(row=row_idx, column=6, value=cpc)
        ws.cell(row=row_idx, column=7, value=cpa)

    total_row = len(rows) + 2
    total_cpa = round(total_budget / total_leads) if total_leads else ""
    totals = ["Итого", "100%", total_budget, total_clicks or "", total_leads or "", "", total_cpa]
    for col, val in enumerate(totals, 1):
        cell = ws.cell(row=total_row, column=col, value=val)
        cell.font = bold

    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 10
    for col in "CDEFG":
        ws.column_dimensions[col].width = 16

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
